import numpy as np
import pandas as pd

from docx.shared import Inches
from docx.shared import Pt
import matplotlib.pyplot as plt
import matplotlib

import io
import pandas as pd
import functions

font = {'family' : 'normal',
        'weight' : 'bold',
        'size'   : 10}

matplotlib.rc('font', **font)

laplas_dict = {
    0:    [        0,   3983,7926, 11791,15542,19146,22575,25804,28814,31594,34134,36433,38493,40320,41924,43319,44520,45543,46407,47128,47725,48214,48610,48928,49180,49379,49534,49653,49744,49813,],
    1:    [        399, 4380,8317, 12172,15910,19497,22907,26115,29103,31859,34375,36650,38686,40490,42073,43448,44630,45637,46485,47193,47778,48257,48645,48956,49202,49396,49547,49664,49752,49819,],
    2:    [        798, 4776,8706, 12552,16276,19847,23237,26424,29389,32121,34614,36864,38877,40658,42220,43574,44738,45728,46562,47257,47831,48300,48679,48983,49224,49413,49560,49674,49760,49825,],
    3:    [        1197,5172,9095, 12930,16640,20194,23565,26730,29673,32381,34850,37076,39065,40824,42364,43699,44845,45818,46638,47320,47882,48341,48713,49010,49245,49430,49573,49683,49767,49831,],
    4:    [        1595,5567,9483, 13307,17003,20540,23891,27035,29955,32639,35083,37286,39251,40988,42507,43822,44950,45907,46712,47381,47932,48382,48745,49036,49266,49446,49585,49693,49774,49836,],
    5:    [        1994,5962,9871, 13683,17364,20884,24215,27337,30234,32894,35314,37493,39435,41149,42647,43943,45053,45994,46784,47441,47982,48422,48778,49061,49286,49461,49598,49702,49781,49841,],
    6:    [        2392,6356,10257,14058,17724,21226,24537,27637,30511,33147,35543,37698,39617,41308,42786,44062,45154,46080,46856,47500,48030,48461,48809,49086,49305,49477,49609,49711,49788,49846,],
    7:    [        2790,6749,10642,14431,18082,21566,24857,27935,30785,33398,35769,37900,39796,41466,42922,44179,45254,46164,46926,47558,48077,48500,48840,49111,49324,49492,49621,49720,49795,49851,],
    8:    [        3188,7142,11026,14803,18439,21904,25175,28230,31057,33646,35993,38100,39973,41621,43056,44295,45352,46246,46995,47615,48124,48537,48870,49134,49343,49506,49632,49728,49801,49856,],
    9:    [        3586,7535,11409,15173,18793,22240,25490,28524,31327,33891,36214,38298,40147,41774,43189,44408,45449,46327,47062,47670,48169,48574,48899,49158,49361,49520,49643,49736,49807,49861,]
}

# xi2_099 = [
#     6.635,9.210,11.345,13.277,15.086,16.812,18.475,20.090,21.666,23.209,24.725,26.217,27.688,29.141,30.578,23.000,
#     33.409,34.805,36.191,37.566,38.932,40.289,41.638,42.980,44.314,45.642,46.963,48.278,49.588,50.892,      
# ]

xi2_099 = [
    3.841, 5.991, 7.815, 9.488, 11.070, 12.592, 14.067, 15.507, 16.919, 18.307, 19.675 
]

# xi2_099 = [6.314, 2.920, 2.353, 2.132, 2.015, 1.943, 1.895, 1.868, 1.833, 1.813, 1.782, 1.761]
xi2_0025 = [
 0.001, 0.051, 0.216, 0.484, 0.831, 1.237,  1.69, 2.18,   2.7, 3.247, 3.816, 4.404, 5.009, 5.629, 6.262, 
 6.908, 7.564, 8.231, 8.907, 9.591, 10.28, 10.98, 11.69, 12.4, 13.12, 13.84, 14.57, 15.31, 16.05, 16.79]
xi2_0975 = [
 5.024, 7.378, 9.348, 11.14, 12.83, 14.45, 16.01, 17.54, 19.02, 20.48, 21.92, 23.34, 24.74, 26.12, 
 27.49, 28.85, 30.19, 31.53, 32.85, 34.17, 35.48, 36.78, 38.08, 39.36, 40.65, 41.92, 43.19, 44.46,]

vars_df = pd.read_excel('Varianty_LR_po_matstat.xlsx') # название файла с вариантами
available_vars = []
for el in list(vars_df.columns.values):
    if isinstance(el, int) or isinstance(el, str) and el.isdigit():
        available_vars.append(el)

available_vars = [str(el) for el in available_vars] + [int(el) for el in available_vars]

def get_variant(variant=0):
    global vars_df
    # if num != 0:
    #     variant = int(input("""Введите номер варианта из таблицы:
    # -> """))
    # else:
    #     variant = num
    
    result_file_name = 'отчет '+str(variant)+'.docx'
    a = k = -100

    a = list(vars_df[variant].iloc[1:].dropna())
    k = int(vars_df[variant].iloc[0])
    return result_file_name, a, k

def pre_return(result_file_name, a, k, document, color, plt=plt):

    h = round((max(a) - min(a)) / k, 2)
    intervals, intervals_middles, absolute_frequency = pre_params(a, k, h)
    other_work = pre_distribution(a, k,h, intervals, intervals_middles, absolute_frequency)


    for el in other_work:
        represent_res(document, el)

    distribution, distribution_clean = {}, {}
    distribution.update({min(a)-0.85*0.1*(max(a) - min(a)) if min(a)>0 else min(a)-1.15*0.1*(max(a) - min(a)) : 0})
    distribution_clean.update({min(a)-1 : 0})
    common_b = 0
    for ii, _ in enumerate(intervals[:-1]):
        common_b += absolute_frequency[ii]
        distribution.update({intervals[ii] : common_b})
        distribution_clean.update({intervals[ii] : absolute_frequency[ii]})

    fig = plt.figure()
    fig.patch.set_facecolor((247/255, 247/255, 247/255))
    
    ax = fig.add_subplot(111)
    ax.set_facecolor((247/255, 247/255, 247/255))

    # plt.margins(0)
    plt.tight_layout(pad=1.2)
    plt.tight_layout()


    x, y = list(distribution_clean.keys())[1:], list(distribution_clean.values())[1:]
    _ = plt.xticks(sorted(intervals_middles), rotation=20)
    _ = plt.step(x, y, color=color, where='post', linewidth=3, alpha=0.8)
    _ = plt.hlines(y[-1], xmin=x[-1], xmax=x[-1]+h, color=color, linewidth=3)
    _ = plt.vlines(x[-1]+h, ymin=0, ymax=y[-1], color=color, linewidth=3)
    _ = plt.ylabel('Абсолютная частота mi')
    for ii, el in enumerate(x): 
        _ = plt.vlines(el, ymax=y[ii], ymin=0, color=color, linewidth=3)
    _ = plt.grid(axis='y')

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=600,) # facecolor=fig.get_facecolor()
    
    return buf, document, intervals, absolute_frequency

def get_distr_type():
    distr_type = input("""Введите номер распределения:
        1. Показательное
        2. Равномерное
        3. Нормальное
    -> """)
    distr_type = int (distr_type) if distr_type.isdigit() else ''
    return distr_type

def post_return(document, result_file_name, distr_type, a, intervals, k, absolute_frequency):

    if distr_type == 0:
        other_work = e_distribution(a, intervals, absolute_frequency)
    elif distr_type == 2:
        other_work = n_distriburion(a, intervals, absolute_frequency)
    elif distr_type == 1:
        other_work = r_distriburion(a, intervals, k, absolute_frequency)
    else:
        other_work = []
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    for el in other_work:
        represent_res(document, el)
    document.add_page_break()
    for el in proove_intervals(a, k):
        represent_res(document, el)
    # document.add_page_break()

    return document
    # document.save(result_file_name)
    # print('В процессе ...')
    # print('Готово ...')
    # print(f"Создан файл с именем '{result_file_name}'")


def laplas(x):
    if x == -999:
        return -0.5
    if x == 999:
        return 0.5
    str_x = str(abs(x)) + '000000'
    i = int(str_x[3])
    j = int(float(str_x[:3])*10)
    minus = int(x / abs(x))
    if i < len(laplas_dict):
        if j < len(laplas_dict[i]):
            return minus * laplas_dict[i][j]/100000
    return 0

def pre_params(a, k, h):
    intervals =[]
    m = min(a) - h
    for _ in list(range(int(str(k).split('.')[0])+1+1))[:-1]:
        m = m + h
        intervals.append(round(m, 2))
    intervals_middles = []
    for ii, el in enumerate(intervals[:-1]):
        intervals_middles.append(round((el + intervals[ii+1]) / 2, 2))
    absolute_frequency = []
    for m in range(len(intervals[:-1])):
        cond = (intervals[m], intervals[m+1])
        count_c = 0 if m != len(intervals[:-1]) else 1
        for el in a:
            if el<=cond[1] and el>=cond[0] if m == 0 else el<=cond[1] and el>cond[0]:
                count_c += 1
        absolute_frequency.append(count_c)
    return intervals, intervals_middles, absolute_frequency


def pre_distribution(a,k,h, intervals,intervals_middles, absolute_frequency, file_name=''):
    intervals_brackets = []
    for i, el in enumerate(intervals[:-1]):
        s = '[' if i == 0 else '('
        s += str(el) + '; ' + str(intervals[i+1]) + ']'
        intervals_brackets.append(s)
        
    other_work = pd.DataFrame(intervals_brackets, columns=['Интервал'])
    other_work['Середина интервала xi*'] = intervals_middles
    other_work['Абсолютная частота mi'] = absolute_frequency
    other_work['mi * xi*'] = other_work['Середина интервала xi*'] * other_work['Абсолютная частота mi']
    x_s_palochkoy = other_work['mi * xi*'].sum() / len(a)
    other_work['Xi* –  x̅'] = (other_work['Середина интервала xi*'] - x_s_palochkoy).apply(lambda x: round(x, 3))
    other_work['(xi –  x̅) ^ 2'] = other_work['Xi* –  x̅'].pow(2).apply(lambda x: round(x, 3))
    other_work['mi * (xi* –  x̅) ^ 2'] = other_work['Абсолютная частота mi'] * other_work['(xi –  x̅) ^ 2']
    s2 = other_work['mi * (xi* –  x̅) ^ 2'].sum() / (len(a) -1)
    other_work.index += 1
    other_work['№'] = other_work.index
    other_work = other_work.set_index('№')
    if file_name:
        other_work.to_excel('первая часть'+file_name)
    
    s = ''
    for e in a:
        s = s + str(e) + ', '
    description = f"""pДана выборка объема n={len(a)}
{s[:-2]}

Xmin = {min(a)}
Xmax = {max(a)}
Значения изучаемой случайной величины (СВ) расположены на отрезке [{min(a)}, {max(a)}]
Разбиваем этот отрезок на k = {k}

h = (Xmax – Xmin) / k = ({max(a)} – {min(a)}) / {k} = {h}"""
    disp = other_work['mi * (xi* –  x̅) ^ 2'].sum()/ (len(a) -1)
    description2 = f"""p
Получены следующие характеристики:
    1. Выборочная средняя (оценка математического ожидания)   
        x̅   = Σ (mi * xi) / n = {round(other_work['mi * xi*'].sum(), 2)} / {len(a)} = 
             = {round(np.mean(a), 2)}
    2. Несмещенная оценка дисперсии (исправленная дисперсия)  
        S^2 = Σ (mi * (xi –  x̅) ^ 2) / (n-1) = {round(other_work['mi * (xi* –  x̅) ^ 2'].sum(), 2)} / {len(a) -1}  =  
             = {round(disp, 2)}
    3. Выборочное среднее квадратичное отклонение (выборочный стандарт) 
        S   = √ (S) =  √ ({round(disp, 2)}) = 
             = {round(disp**0.5, 2)}"""
    description3 = """h

Построим гистограмму частот mi или гистограмму относительных частот mi/n :"""
    return 'hАнализ представленной выборки:', description, other_work, description2, description3


def e_distribution(a, intervals, absolute_frequency, file_name=''):
    x_s_palochkoy = np.mean(a)
    other_work = pd.DataFrame(intervals[:-1], columns=['xi'])
    other_work['xi+1'] = intervals[1:]
    other_work['mi'] = absolute_frequency
    other_work['xi*'] = (other_work['xi']+ other_work['xi+1']) / 2
    other_work['xi* * mi'] = other_work['xi*'] * other_work['mi']
    other_work['(xi*)^2 * mi'] = other_work['xi*'].pow(2) * other_work['mi']
    other_work['(xi* -  x̅)^2 * mi'] = (other_work['xi*'] - x_s_palochkoy).pow(2) * other_work['mi']
    e = 2.718281828459045235360287471352 
    other_work['e^(-лямбда * xi)'] = np.power([e]*other_work.shape[0], list(-1 * other_work['xi'] / x_s_palochkoy))
    other_work['e^(-лямбда * xi+1)'] = np.power([e]*other_work.shape[0], list(-1 * other_work['xi+1'] / x_s_palochkoy))
    other_work['pi'] = other_work['e^(-лямбда * xi)'] - other_work['e^(-лямбда * xi+1)']
    other_work['miT'] = len(a) * other_work['pi']
    other_work['mi - miT'] = other_work['mi'] - other_work['miT']
    other_work['(mi - miT)^2 / miT'] = other_work['mi - miT'] * other_work['mi - miT'] / other_work['miT']
    other_work.index +=1
    other_work['№'] = other_work.index
    other_work = other_work.set_index('№')
    if file_name:
        other_work.to_excel('показательное распределние'+file_name)
    
    r_ = len(other_work)-3
    w_t = xi2_099[r_-1]
    w_ = round(other_work['(mi - miT)^2 / miT'].sum(), 2)
    description = f"""
Σ mi = {sum(other_work['mi'])}
Σ miT = {round(sum(other_work['miT']), 2)}
Σ (xi* * mi) = {round(sum(other_work['xi* * mi']), 2)}
Σ ((xi*)^2 * mi) = {round(sum(other_work['(xi*)^2 * mi']), 2)}
Σ ((xi* -  x̅)^2 * mi) = {round(sum(other_work['(xi* -  x̅)^2 * mi']))}
Σ pi = {round(sum(other_work['pi']), 2)}
Σ ((mi - miT)^2 / miT) = {round(sum(other_work['(mi - miT)^2 / miT']), 2)}"""
    description2 = f"""r = n - 2 - 1 = {other_work.shape[0]} - 3 = {r_}
X^2 набл = Σ ((mi - miT)^2 / miT) = {round(sum(other_work['(mi - miT)^2 / miT']), 2)}
X^2 кр = {w_t}

X^2 набл {'<=' if w_ <= w_t else '>'} X^2 кр 
{w_} {'<=' if w_ <= w_t else '>'} {w_t}
где X^2 кр берется из таблицы квантилей X^2 распределения

Гипотеза {'НЕ ' if abs(w_) <= w_t else ''}отвергается на уровне значимости α=0,05"""
    return 'hПоказательное распределение', other_work, 'p'+description, 'p'+description2


def n_distriburion(a, intervals, absolute_frequency, file_name=''):
    other_work = pd.DataFrame(intervals[:-1], columns=['xi'])
    other_work['xi+1'] = intervals[1:]
    other_work['xi*'] = (other_work['xi'] + other_work['xi+1']) / 2
    other_work['mi'] = absolute_frequency
    other_work['xi* * mi'] = other_work['xi*'] * other_work['mi']
    xi_mean = sum(other_work['xi* * mi']) / len(a)
    other_work['(xi* -  x̅)^2 * mi'] = (other_work['xi*'] - xi_mean) ** 2 * other_work['mi']
    other_work['zi'] = (other_work['xi'] - xi_mean) / np.std(a)
    other_work['zi'].iloc[0] = -999
    other_work['zi+1'] = (other_work['xi+1'] - xi_mean) / np.std(a)
    other_work['zi+1'].iloc[-1] = 999
    other_work['Фzi'] = other_work['zi'].apply(laplas)
    other_work['Фzi+1'] = other_work['zi+1'].apply(laplas)
    other_work['pi'] = other_work['Фzi+1'] - other_work['Фzi']
    other_work['miT'] = len(a) * other_work['pi']
    other_work['mi - miT'] = other_work['mi'] - other_work['miT']
    other_work['(mi - miT)^2 / miT'] = other_work['mi - miT'].pow(2) / other_work['miT'] 
    other_work.index +=1
    other_work['№'] = other_work.index
    other_work = other_work.set_index('№')
    if file_name:
        other_work.to_excel('нормальное распределение'+file_name)
        
    r_ = other_work.shape[0] -3
    w_t = xi2_099[r_-1]
    w_ = round(sum(other_work['(mi - miT)^2 / miT']), 2) 
    description = f"""
Σ xi* = {round(sum(other_work['xi*']), 2)}
Σ mi = {sum(other_work['mi'])}
Σ xi* * mi = {round(sum(other_work['xi* * mi']), 2)}
Σ (xi* -  x̅)^2 * mi = {round(sum(other_work['(xi* -  x̅)^2 * mi'])/other_work.shape[0] ,2)}
Σ pi = {sum(other_work['pi'])}
Σ mi - miT = {round(sum(other_work['mi - miT']), 2)}
W = Σ (mi - miT)^2 / miT = {round(sum(other_work['(mi - miT)^2 / miT']), 2)}"""
    description2 = f"""r = n - 2 -1 = {other_work.shape[0]} - 3 = {r_}
X^2 набл = Σ((mi - miT)^2 / miT) = {round(sum(other_work['(mi - miT)^2 / miT']), 2)}
X^2 кр = {w_t}

X^2 набл {'<=' if w_ <= w_t else '>'} 'X^2 кр
{w_} {'<=' if w_ <= w_t else '>'} {w_t}
где X^2 кр берется из таблицы квантилей X^2 распределения"""

    description3 = """hГипотеза {'НЕ ' if abs(w_) <= w_t else ''}отвергается на уровне значимости α=0,05"""

    return 'hНормальное распределение', other_work, 'p'+description, 'p'+description2, description3



def r_distriburion(a, intervals,k, absolute_frequency, file_name=''):
    other_work = pd.DataFrame(intervals[:-1], columns=['xi'])
    other_work['xi+1'] = intervals[1:]
    other_work['xi*'] = (other_work['xi']+ other_work['xi+1']) / 2
    other_work['mi'] = absolute_frequency
    other_work['mi * x*'] = other_work['mi'] * other_work['xi*']
    xi_mean = other_work['xi*'].mean()
    other_work['xi* -  x̅'] = other_work['xi*'] - xi_mean
    other_work['(xi* -  x̅)^2'] = other_work['xi* -  x̅'].pow(2)
    disp = np.std(a)**2
    other_work['pi'] = round((max(a) - min(a)) / k, 2) / (2* disp**0.5 * 3**0.5)
    other_work['pi'].iloc[0] = (other_work['xi*'].iloc[1] - xi_mean + disp**0.5 * 3**0.5) / (2* disp**0.5 * 3**0.5)
    other_work['pi'].iloc[-1] = (xi_mean + disp**0.5 * 3**0.5 - other_work['xi*'].iloc[-2]) / (2* disp**0.5 * 3**0.5)
    other_work['miT'] = len(a) * other_work['pi']
    other_work['mi - miT'] = other_work['mi'] - other_work['miT']
    other_work['(mi - miT)'] = other_work['mi - miT'].pow(2)
    other_work['W'] = other_work['(mi - miT)'] / other_work['miT']
    other_work.index +=1
    other_work['№'] = other_work.index
    other_work = other_work.set_index('№')
    if file_name:
        other_work.to_excel('равномерное распределение'+file_name)
    r_ = len(other_work) -3
    w_t = xi2_099[r_-1]
    w_ = round(sum(other_work['W']), 2)
    description = f"""
Σ mi* = {len(a)}
Σ pi* = 1
Σ mi* = {len(a)}
Xнабл^2 = Σ W = {round(other_work['W'].sum(), 2)}"""
    description2 = f"""r = n - 2 -1 = {other_work.shape[0]} - 2 -1 = {r_}
X^2 кр = {w_t}

X^2 набл {'<=' if w_ <= w_t else '>'} X^2 кр
{w_} {'<=' if w_ <= w_t else '>'} {w_t}
где X^2 кр берется из таблицы квантилей X^2 распределения

Гипотеза {'НЕ ' if abs(w_) <= w_t else ''}отвергается на уровне значимости α=0,05"""
    return 'hРавномерное распределение', other_work, 'p'+description, 'p'+description2



def proove_intervals(a, k):
    disp = np.std(a)**2
    x_1 = np.mean(a) - disp**0.5 * xi2_099[k-2] / k**0.5
    x_2 = np.mean(a) + disp**0.5 * xi2_099[k-2] / k**0.5
    description = f"""pДоверительный интервал для среднего значения при неизвестной дисперсии
    С доверительной вероятностью  (надежностью) (1-0.05) среднеe значениe накрывается интервалом
     x̅ - S / √ (n) * t < a <  x̅ + S / √ (n) * t
    {round(np.mean(a), 2)} - {round(disp**0.5, 2)} / {round(k**0.5, 2)} * {round(xi2_099[k-2], 2)} < a < {round(np.mean(a), 2)} + {round(disp**0.5, 2)} / {round(k**0.5, 2)} * {round(xi2_099[k-2], 2)}
    {round(x_1, 2)} < a < {round(x_2, 2)}

Доверительный интервал для дисперсии при неизвестном математическом ожидании
    С доверительной вероятностью  (надежностью) (1-0.05) неизвестная дисперсия σ^2 накрывается интервалом
    (n-1) * S^2 / X(а/2),n-1 < σ^2 < (n-1) * S^2 / X(1-а/2),n-1
    {k-1} * {round(disp, 2)} / {xi2_0975[k-1]} < σ^2 < {k-1} * {round(disp, 2)} / {xi2_0025[k-1]}
    {round((k-1) * disp / xi2_0975[k-2], 2)} < σ^2 < {round((k-1) * disp / xi2_0025[k-2] ,2)}"""
    return 'hДоверительные интервалы', description

def represent_res(document, x):

    if isinstance(x, pd.core.frame.DataFrame):
        records = list(x.values)
        columns = list(x.columns)
        table = document.add_table(rows=1, cols=len(columns))
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        for i, el in enumerate(columns):
            hdr_cells[i].text = el   
        for record in records:
            row_cells = table.add_row().cells
            for i, el in enumerate(record):
                row_cells[i].text = str(round(el, 4)) if str(el)[4:].isdigit() else str(el)[:11]
            
    elif isinstance(x, str):
        if x[0] == 'h':
            document.add_paragraph('').add_run(x[1:]).bold = True
#             print('headline', x[1:])
        
        elif x[0] == 'p':
            document.add_paragraph(x[1:])
#             print('paragraph', x[1:])
    else:
        print('hz', x[1:])